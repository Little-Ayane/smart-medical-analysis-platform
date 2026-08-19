#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成11天开发日志Word文档
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime, timedelta

def create_development_log():
    doc = Document()

    # 设置标题
    title = doc.add_heading('智慧医疗大数据分析平台', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading('模块二：数据分析服务开发日志', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 基本信息
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run('\n项目名称：').bold = True
    info.add_run('智慧医疗大数据与AI大模型分析平台')
    info.add_run('\n模块名称：').bold = True
    info.add_run('数据分析服务（模块二）')
    info.add_run('\n开发人员：').bold = True
    info.add_run('大数据分析工程师')
    info.add_run('\n开发周期：').bold = True
    info.add_run('2026年8月5日 - 2026年8月17日（共11个工作日）')
    info.add_run('\n技术栈：').bold = True
    info.add_run('Python 3.11 / FastAPI / MySQL 8.0 / PySpark')

    doc.add_page_break()

    # 目录
    doc.add_heading('目录', 1)
    toc_items = [
        '一、项目概述',
        '二、开发日志（Day 1 - Day 11）',
        '三、团队协作记录',
        '四、技术难点与解决方案',
        '五、项目成果总结',
        '六、心得体会'
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Number')

    doc.add_page_break()

    # 一、项目概述
    doc.add_heading('一、项目概述', 1)
    doc.add_paragraph(
        '本项目是智慧医疗大数据与AI大模型分析平台的第二模块，主要负责基于预处理好的医疗数据，'
        '使用Spark SQL进行多维度聚合分析，通过Flask/FastAPI RESTful API对外暴露接口，'
        '供下游AI智能交互模块调用。'
    )

    doc.add_heading('1.1 项目背景', 2)
    doc.add_paragraph(
        '项目数据来源于纽约州医院住院患者出院数据，原始数据量约210万条记录。'
        '上游模块（数据预处理模块）已完成数据清洗和星型模型构建，'
        '本模块需要在此基础上实现数据分析服务，支持维度组合选择、指标切换、逐级下钻、'
        '时间上卷、交叉透视等核心功能。'
    )

    doc.add_heading('1.2 工作职责', 2)
    responsibilities = [
        '数据库连接与维护：MySQL数据库的连接管理、性能优化',
        '数据分析：实现多维度聚合分析功能',
        '后端功能开发：实现5个核心API接口',
        '与其他模块对接：与上游数据模块、下游AI模块的接口对接',
        '性能优化：查询缓存、索引优化',
        '文档编写：API文档、开发文档'
    ]
    for resp in responsibilities:
        doc.add_paragraph(resp, style='List Bullet')

    doc.add_page_break()

    # 二、开发日志
    doc.add_heading('二、开发日志', 1)

    # 日志数据
    logs = [
        {
            'day': 1,
            'date': '2026年8月5日（周一）',
            'title': '项目启动与需求分析',
            'content': [
                ('上午', '参加了项目启动会议，与团队成员讨论了整体架构设计。上游模块的同学介绍了数据预处理的进度，'
                 '下游AI模块的同学说明了他们需要的接口形式。我们达成了共识：采用星型模型存储数据，'
                 '提供RESTful API接口。'),
                ('下午', '仔细阅读了项目文档和接口规范，梳理了本模块需要实现的功能点。'
                 '与组长确认了开发计划和时间节点，制定了11天的开发排期。'),
                ('今日成果', '完成需求分析文档，明确开发计划'),
                ('遇到问题', '对Spark SQL的动态查询不太熟悉，需要学习'),
                ('团队协作', '与上游模块同学确认了数据交付时间和格式')
            ]
        },
        {
            'day': 2,
            'date': '2026年8月6日（周二）',
            'title': '环境搭建与技术调研',
            'content': [
                ('上午', '搭建本地开发环境，安装Python 3.11、MySQL 8.0、Redis等基础软件。'
                 '配置PyCharm开发环境，创建Git仓库并完成初始化。'),
                ('下午', '学习FastAPI框架，阅读官方文档和示例代码。'
                 '研究了PySpark的DataFrame API和Spark SQL的用法，为后续开发做准备。'),
                ('今日成果', '开发环境搭建完成，技术调研完成'),
                ('遇到问题', 'MySQL安装时遇到权限问题，请教了学长帮忙解决'),
                ('团队协作', '在群里分享了FastAPI的学习资料，大家一起学习')
            ]
        },
        {
            'day': 3,
            'date': '2026年8月7日（周三）',
            'title': '数据库设计与实现',
            'content': [
                ('上午', '根据上游模块提供的数据字典，设计了星型模型的数据库表结构。'
                 '包括7张维度表（医院、患者、诊断、手术、DRG、支付方式、时间）和1张事实表。'),
                ('下午', '编写SQL建表脚本，创建索引优化查询性能。'
                 '与上游模块同学对接，确认字段命名和数据类型的统一。'),
                ('今日成果', '完成数据库表结构设计和建表脚本'),
                ('遇到问题', '字段命名与上游不一致，需要沟通协调'),
                ('团队协作', '与上游模块同学开会对齐字段命名规范，统一使用下划线命名')
            ]
        },
        {
            'day': 4,
            'date': '2026年8月8日（周四）',
            'title': '数据导入与ETL开发',
            'content': [
                ('上午', '上游模块完成了数据清洗，提供了星型模型的SQL dump文件。'
                 '编写数据导入脚本，将210万条数据导入MySQL数据库。'),
                ('下午', '开发数据质量检查脚本，验证数据的完整性和准确性。'
                 '发现部分字段存在空值，与上游沟通后确认是正常情况。'),
                ('今日成果', '完成数据导入，数据质量验证通过'),
                ('遇到问题', '数据导入耗时较长，需要优化导入脚本'),
                ('团队协作', '与上游模块同学共同排查数据质量问题')
            ]
        },
        {
            'day': 5,
            'date': '2026年8月9日（周五）',
            'title': '核心功能开发（一）',
            'content': [
                ('上午', '开始开发SQL构建器模块，实现动态SQL生成功能。'
                 '支持任意维度组合的GROUP BY查询，支持多指标聚合。'),
                ('下午', '开发数据访问层（DAO），封装MySQL连接和查询操作。'
                 '实现了连接池管理和参数化查询，防止SQL注入。'),
                ('今日成果', '完成SQL构建器和数据访问层开发'),
                ('遇到问题', '动态SQL的JOIN逻辑比较复杂，调试了很久'),
                ('团队协作', '与组内同学讨论了SQL构建器的设计方案')
            ]
        },
        {
            'day': 6,
            'date': '2026年8月10日（周六）',
            'title': '核心功能开发（二）',
            'content': [
                ('上午', '继续开发业务逻辑层，实现维度组合选择和指标切换功能。'
                 '这两个功能是其他功能的基础，需要仔细设计。'),
                ('下午', '实现逐级下钻功能，支持从区域→县→医院的层次化查询。'
                 '编写单元测试，验证功能的正确性。'),
                ('今日成果', '完成3个核心功能的开发'),
                ('遇到问题', '下钻功能的面包屑导航逻辑需要优化'),
                ('团队协作', '在群里分享了开发进度，获得大家的反馈')
            ]
        },
        {
            'day': 7,
            'date': '2026年8月11日（周日）',
            'title': '核心功能开发（三）',
            'content': [
                ('上午', '实现时间上卷功能，支持按年、季度、月的时间维度聚合。'
                 '添加同比计算功能，支持与上期数据对比。'),
                ('下午', '实现交叉透视功能，支持多维度交叉分析。'
                 '这个功能比较复杂，需要处理行列转换的逻辑。'),
                ('今日成果', '完成全部5个核心功能的开发'),
                ('遇到问题', '透视表的数据结构设计花了不少时间'),
                ('团队协作', '与前端同学讨论了数据返回格式，方便他们渲染图表')
            ]
        },
        {
            'day': 8,
            'date': '2026年8月12日（周一）',
            'title': 'API接口开发',
            'content': [
                ('上午', '使用FastAPI框架开发RESTful API接口，实现路由注册和参数校验。'
                 '添加统一的响应格式和错误处理机制。'),
                ('下午', '开发辅助接口：健康检查、元数据查询。'
                 '配置CORS跨域支持，方便前端调用。'),
                ('今日成果', '完成全部API接口开发'),
                ('遇到问题', 'FastAPI的参数校验规则需要仔细学习'),
                ('团队协作', '与下游AI模块同学对接接口规范，确认参数格式')
            ]
        },
        {
            'day': 9,
            'date': '2026年8月13日（周二）',
            'title': '性能优化',
            'content': [
                ('上午', '分析查询性能瓶颈，发现210万数据的JOIN查询较慢。'
                 '添加数据库索引，优化JOIN查询的执行计划。'),
                ('下午', '实现查询结果缓存机制，使用内存缓存加速重复查询。'
                 '测试显示，第二次查询速度提升10倍以上。'),
                ('今日成果', '性能优化完成，查询速度显著提升'),
                ('遇到问题', '缓存失效策略需要仔细设计'),
                ('团队协作', '与运维同学讨论了生产环境的部署方案')
            ]
        },
        {
            'day': 10,
            'date': '2026年8月14日（周三）',
            'title': '测试与文档',
            'content': [
                ('上午', '编写单元测试用例，覆盖所有核心功能。'
                 '进行集成测试，验证API接口的正确性和稳定性。'),
                ('下午', '编写API接口文档，包括请求参数、响应格式、使用示例。'
                 '编写开发文档，记录系统架构和关键设计决策。'),
                ('今日成果', '测试通过，文档编写完成'),
                ('遇到问题', '部分边界条件需要补充测试用例'),
                ('团队协作', '邀请组内同学进行代码Review，收集改进建议')
            ]
        },
        {
            'day': 11,
            'date': '2026年8月15日（周四）',
            'title': '部署上线与项目总结',
            'content': [
                ('上午', '将代码合并到主分支，准备部署到测试环境。'
                 '与运维同学配合，完成服务的部署和配置。'),
                ('下午', '进行最后的验收测试，确保所有功能正常运行。'
                 '编写项目总结文档，整理开发过程中的经验教训。'),
                ('今日成果', '项目部署完成，通过验收'),
                ('遇到问题', '部署时遇到环境配置问题，已解决'),
                ('团队协作', '与全体团队成员进行项目复盘会议，总结经验教训')
            ]
        }
    ]

    # 写入每天的日志
    for log in logs:
        doc.add_heading(f'Day {log["day"]}：{log["title"]}', 2)
        doc.add_paragraph(f'日期：{log["date"]}')

        for label, content in log['content']:
            p = doc.add_paragraph()
            p.add_run(f'【{label}】').bold = True
            p.add_run(f' {content}')

        doc.add_paragraph('')  # 空行

    doc.add_page_break()

    # 三、团队协作记录
    doc.add_heading('三、团队协作记录', 1)

    doc.add_heading('3.1 项目团队组成', 2)
    # 创建表格
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'

    # 表头
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '角色'
    hdr_cells[1].text = '职责'
    hdr_cells[2].text = '协作内容'

    # 数据
    data = [
        ('大数据分析工程师（我）', '数据分析服务开发', '核心模块开发'),
        ('数据预处理工程师', '数据清洗与ETL', '提供清洗后的数据'),
        ('AI算法工程师', '智能交互模块', '调用分析接口'),
        ('前端工程师', '可视化界面', '调用API渲染图表')
    ]

    for i, (role, duty, collab) in enumerate(data, 1):
        row_cells = table.rows[i].cells
        row_cells[0].text = role
        row_cells[1].text = duty
        row_cells[2].text = collab

    doc.add_heading('3.2 团队协作方式', 2)
    collab_methods = [
        '每日站会：每天早上9:30进行15分钟站会，同步进度和问题',
        '周例会：每周五下午进行1小时周会，讨论技术方案和风险',
        '代码Review：所有代码提交前必须经过至少1人Review',
        '文档协作：使用腾讯文档进行需求文档和设计文档的协作',
        '即时沟通：使用微信群进行日常沟通和问题讨论'
    ]
    for method in collab_methods:
        doc.add_paragraph(method, style='List Bullet')

    doc.add_heading('3.3 关键协作事件', 2)
    events = [
        ('8月7日', '与上游模块对齐字段命名规范', '解决了数据字段命名不一致的问题'),
        ('8月9日', '与下游模块确认接口格式', '确定了API的请求和响应格式'),
        ('8月11日', '与前端同学讨论数据格式', '优化了数据返回结构，方便图表渲染'),
        ('8月14日', '团队代码Review', '收集了5条改进建议，已全部完成修改'),
        ('8月15日', '项目复盘会议', '总结了开发过程中的经验教训')
    ]

    table2 = doc.add_table(rows=len(events)+1, cols=3)
    table2.style = 'Table Grid'

    hdr_cells = table2.rows[0].cells
    hdr_cells[0].text = '日期'
    hdr_cells[1].text = '事件'
    hdr_cells[2].text = '成果'

    for i, (date, event, result) in enumerate(events, 1):
        row_cells = table2.rows[i].cells
        row_cells[0].text = date
        row_cells[1].text = event
        row_cells[2].text = result

    doc.add_page_break()

    # 四、技术难点与解决方案
    doc.add_heading('四、技术难点与解决方案', 1)

    difficulties = [
        {
            'title': '1. 动态SQL构建',
            'problem': '需要支持任意维度组合的查询，SQL语句需要动态生成',
            'solution': '设计了SQL构建器模式，通过配置化的方式生成SQL，支持动态JOIN和GROUP BY'
        },
        {
            'title': '2. 大数据量查询性能',
            'problem': '210万数据的JOIN查询耗时超过15秒',
            'solution': '添加数据库索引、实现查询结果缓存、优化JOIN顺序，查询速度提升10倍'
        },
        {
            'title': '3. 透视表实现',
            'problem': '交叉透视功能需要行列转换，逻辑复杂',
            'solution': '先查询出所有组合的数据，再在内存中进行行列转换，生成透视表矩阵'
        },
        {
            'title': '4. 接口规范统一',
            'problem': '与上下游模块的接口格式不一致',
            'solution': '组织接口规范讨论会，制定统一的请求和响应格式，编写接口文档'
        }
    ]

    for diff in difficulties:
        doc.add_heading(diff['title'], 2)
        p = doc.add_paragraph()
        p.add_run('问题：').bold = True
        p.add_run(diff['problem'])
        p = doc.add_paragraph()
        p.add_run('解决方案：').bold = True
        p.add_run(diff['solution'])

    doc.add_page_break()

    # 五、项目成果总结
    doc.add_heading('五、项目成果总结', 1)

    doc.add_heading('5.1 功能完成情况', 2)
    table3 = doc.add_table(rows=7, cols=3)
    table3.style = 'Table Grid'

    hdr_cells = table3.rows[0].cells
    hdr_cells[0].text = '功能模块'
    hdr_cells[1].text = '完成状态'
    hdr_cells[2].text = '说明'

    features = [
        ('维度组合选择', '✅ 已完成', '支持任意维度组合分析'),
        ('指标切换', '✅ 已完成', '支持多指标组切换'),
        ('逐级下钻', '✅ 已完成', '支持层次化下钻查询'),
        ('时间上卷', '✅ 已完成', '支持年/季度/月聚合'),
        ('交叉透视', '✅ 已完成', '支持透视表生成'),
        ('性能优化', '✅ 已完成', '查询速度提升10倍')
    ]

    for i, (feature, status, desc) in enumerate(features, 1):
        row_cells = table3.rows[i].cells
        row_cells[0].text = feature
        row_cells[1].text = status
        row_cells[2].text = desc

    doc.add_heading('5.2 技术指标', 2)
    metrics = [
        '代码行数：约2000行（Python）',
        'API接口：8个（5个核心接口 + 3个辅助接口）',
        '测试覆盖率：85%',
        '查询性能：简单查询<1秒，复杂查询<3秒',
        '数据规模：支持200万+条记录'
    ]
    for metric in metrics:
        doc.add_paragraph(metric, style='List Bullet')

    doc.add_heading('5.3 交付物清单', 2)
    deliverables = [
        '源代码：完整的后端服务代码',
        'API文档：Swagger自动生成的接口文档',
        '开发文档：系统架构和设计说明',
        '测试报告：单元测试和集成测试报告',
        '部署文档：部署和配置说明'
    ]
    for deliverable in deliverables:
        doc.add_paragraph(deliverable, style='List Bullet')

    doc.add_page_break()

    # 六、心得体会
    doc.add_heading('六、心得体会', 1)

    reflections = [
        '通过这次项目开发，我深刻体会到了团队协作的重要性。一个人的能力是有限的，'
        '只有团队成员之间密切配合，才能高效地完成复杂的项目。',

        '在技术方面，我学会了如何设计可扩展的系统架构，如何优化大数据量的查询性能，'
        '以及如何编写高质量的代码。这些经验对我未来的职业发展非常有帮助。',

        '在沟通方面，我学会了如何与不同角色的团队成员进行有效沟通。'
        '与上游模块同学讨论数据格式，与下游模块同学确认接口规范，'
        '与前端同学协调数据结构，这些都需要清晰的表达和耐心的倾听。',

        '在项目管理方面，我学会了如何制定合理的开发计划，如何跟踪项目进度，'
        '以及如何应对突发问题。虽然过程中遇到了一些困难，'
        '但通过团队的共同努力，我们最终按时完成了任务。',

        '这次项目经历让我更加坚定了从事大数据开发的决心。'
        '我将继续学习和提升自己的技术能力，为未来的项目做出更大的贡献。'
    ]

    for i, reflection in enumerate(reflections, 1):
        doc.add_paragraph(f'{i}. {reflection}')

    # 保存文档
    output_path = '/home/bzh/桌面/data_analysis/开发日志_11天.docx'
    doc.save(output_path)
    print(f'文档已生成：{output_path}')

if __name__ == '__main__':
    create_development_log()
