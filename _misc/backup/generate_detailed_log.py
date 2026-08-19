#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成详细的11天开发日志Word文档
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def set_font(run, name='宋体', size=12, bold=False):
    """设置字体"""
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)


def add_title(doc, text, level=0):
    """添加标题"""
    heading = doc.add_heading(text, level)
    return heading


def add_paragraph_with_font(doc, text, bold=False, indent=True):
    """添加段落并设置字体"""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)
    run = p.add_run(text)
    set_font(run, size=12, bold=bold)
    return p


def create_detailed_log():
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # ==================== 封面 ====================
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('智慧医疗大数据分析平台')
    set_font(run, name='黑体', size=28, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('模块二：数据分析服务开发日志')
    set_font(run, name='黑体', size=20)

    doc.add_paragraph()
    doc.add_paragraph()

    info_items = [
        ('项目名称', '智慧医疗大数据与AI大模型分析平台'),
        ('模块名称', '数据分析服务（模块二）'),
        ('开发人员', '大数据分析工程师'),
        ('开发周期', '2026年8月5日 - 2026年8月17日（共11个工作日）'),
        ('技术栈', 'Python 3.11 / FastAPI / MySQL 8.0 / PySpark'),
    ]

    for label, value in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'{label}：')
        set_font(run, size=14, bold=True)
        run = p.add_run(value)
        set_font(run, size=14)

    doc.add_page_break()

    # ==================== 目录 ====================
    add_title(doc, '目  录', 1)

    toc_items = [
        '一、项目概述',
        '二、开发环境',
        '三、开发日志',
        '    Day 1  项目启动与需求分析',
        '    Day 2  环境搭建与技术调研',
        '    Day 3  数据库设计与实现',
        '    Day 4  数据导入与ETL开发',
        '    Day 5  核心功能开发（一）',
        '    Day 6  核心功能开发（二）',
        '    Day 7  核心功能开发（三）',
        '    Day 8  API接口开发',
        '    Day 9  性能优化与缓存',
        '    Day 10 测试与文档编写',
        '    Day 11 部署上线与项目总结',
        '四、团队协作记录',
        '五、技术难点与解决方案',
        '六、项目成果总结',
        '七、心得体会与展望'
    ]

    for item in toc_items:
        p = doc.add_paragraph(item)
        set_font(p.runs[0], size=12)

    doc.add_page_break()

    # ==================== 一、项目概述 ====================
    add_title(doc, '一、项目概述', 1)

    add_title(doc, '1.1 项目背景', 2)
    add_paragraph_with_font(doc,
        '本项目是智慧医疗大数据与AI大模型分析平台的重要组成部分。该平台旨在利用大数据技术和人工智能技术，'
        '对医疗数据进行深度分析和挖掘，为医疗决策提供数据支持。平台分为三大模块：'
    )
    add_paragraph_with_font(doc,
        '模块一：数据预处理与持久化（上游）—— 负责原始数据的清洗、转换和存储；'
        '模块二：数据分析服务（本模块）—— 负责基于清洗后的数据进行多维度分析；'
        '模块三：AI智能交互（下游）—— 负责基于分析结果提供智能问答服务。'
    )
    add_paragraph_with_font(doc,
        '本模块作为中间层，承上启下，既要对接上游的数据输出，又要为下游提供稳定高效的API接口。'
        '数据来源为纽约州医院住院患者出院数据，原始数据量约210万条记录，包含患者基本信息、诊断信息、'
        '治疗信息、费用信息等多个维度。'
    )

    add_title(doc, '1.2 工作职责', 2)
    add_paragraph_with_font(doc, '作为大数据分析工程师，我的主要工作职责包括：')

    responsibilities = [
        '数据库连接与维护：负责MySQL数据库的连接管理、性能监控和优化，确保数据库服务稳定运行；',
        '数据分析功能开发：实现多维度聚合分析功能，包括维度组合选择、指标切换、逐级下钻、时间上卷、交叉透视等核心功能；',
        '后端API开发：使用FastAPI框架开发RESTful API接口，提供统一的数据访问接口；',
        '性能优化：通过索引优化、查询缓存等手段，提升大数据量下的查询性能；',
        '团队协作：与上游数据模块、下游AI模块、前端可视化模块进行接口对接和技术协调；',
        '文档编写：编写API接口文档、开发文档、测试报告等技术文档。'
    ]

    for resp in responsibilities:
        p = doc.add_paragraph(resp, style='List Bullet')
        set_font(p.runs[0], size=12)

    add_title(doc, '1.3 技术选型', 2)
    add_paragraph_with_font(doc, '经过技术调研和团队讨论，我们选择了以下技术栈：')

    tech_stack = [
        '编程语言：Python 3.11 —— 语法简洁，生态丰富，适合快速开发；',
        'Web框架：FastAPI —— 高性能异步框架，自动生成API文档，支持类型校验；',
        '数据库：MySQL 8.0 —— 成熟稳定的关系型数据库，支持大数据量存储；',
        '数据处理：PySpark —— 分布式计算框架，适合大数据处理；',
        '缓存：Redis —— 高性能内存数据库，用于查询结果缓存；',
        'ORM：PyMySQL —— 轻量级MySQL驱动，性能好，易于使用。'
    ]

    for tech in tech_stack:
        p = doc.add_paragraph(tech, style='List Bullet')
        set_font(p.runs[0], size=12)

    doc.add_page_break()

    # ==================== 二、开发环境 ====================
    add_title(doc, '二、开发环境', 1)

    add_paragraph_with_font(doc, '本次开发使用的环境配置如下：')

    env_items = [
        '操作系统：Ubuntu 22.04 LTS（虚拟机环境）',
        'Python版本：Python 3.11.0',
        'MySQL版本：MySQL 8.0.30',
        '开发工具：VS Code + PyCharm',
        '版本控制：Git + GitHub',
        '数据库管理：MySQL Workbench',
        '接口测试：Postman / Swagger UI'
    ]

    for item in env_items:
        p = doc.add_paragraph(item, style='List Bullet')
        set_font(p.runs[0], size=12)

    doc.add_page_break()

    # ==================== 三、开发日志 ====================
    add_title(doc, '三、开发日志', 1)

    # ---- Day 1 ----
    add_title(doc, 'Day 1：项目启动与需求分析（2026年8月5日 周一）', 2)

    add_title(doc, '【今日工作内容】', 3)
    add_paragraph_with_font(doc,
        '上午9:00，参加了项目启动会议。会议由项目经理主持，全体团队成员参加。'
        '会议上，项目经理详细介绍了项目的整体背景、目标和各模块的职责划分。'
        '上游模块的同学介绍了数据预处理的进展情况，他们已经完成了原始数据的清洗工作，'
        '正在构建星型模型的维度表和事实表。下游AI模块的同学说明了他们对接口的需求，'
        '希望能够提供统一的、参数化的API接口，方便他们的LangChain Agent调用。'
    )
    add_paragraph_with_font(doc,
        '下午，我仔细阅读了项目文档，包括《项目企划书》《接口规范》《数据字典》等。'
        '通过阅读这些文档，我对项目的整体架构和各模块之间的关系有了清晰的认识。'
        '然后，我梳理了本模块需要实现的功能点，包括：维度组合选择、指标切换、逐级下钻、'
        '时间上卷、交叉透视等5个核心功能，以及健康检查、元数据查询等辅助功能。'
    )
    add_paragraph_with_font(doc,
        '最后，我与组长进行了沟通，确认了开发计划和时间节点。我们决定采用敏捷开发的方式，'
        '将11天的开发周期分为4个阶段：第一阶段（Day 1-4）完成数据层建设，'
        '第二阶段（Day 5-7）完成核心功能开发，第三阶段（Day 8-9）完成API开发和性能优化，'
        '第四阶段（Day 10-11）完成测试、文档和部署。'
    )

    add_title(doc, '【遇到的问题】', 3)
    add_paragraph_with_font(doc,
        '今天主要是学习和理解阶段，没有遇到具体的技术问题。但是对Spark SQL的动态查询不太熟悉，'
        '需要在后续开发中学习和实践。另外，对FastAPI框架的使用也需要进一步学习。'
    )

    add_title(doc, '【团队协作】', 3)
    add_paragraph_with_font(doc,
        '今天与上游模块的同学进行了初步沟通，确认了数据交付的时间和格式。'
        '他们计划在Day 4提供星型模型的SQL dump文件，包含7张维度表和1张事实表。'
        '与下游AI模块的同学讨论了接口的设计原则，他们希望接口能够参数化，'
        '一个接口能够覆盖多种查询场景，而不是为每种场景设计单独的接口。'
    )

    add_title(doc, '【今日成果】', 3)
    add_paragraph_with_font(doc, '1. 完成项目启动会议，明确项目目标和职责划分；')
    add_paragraph_with_font(doc, '2. 阅读项目文档，梳理功能需求清单；')
    add_paragraph_with_font(doc, '3. 制定11天开发计划，明确各阶段任务；')
    add_paragraph_with_font(doc, '4. 与上下游模块完成初步沟通。')

    add_title(doc, '【明日计划】', 3)
    add_paragraph_with_font(doc, '1. 搭建本地开发环境；')
    add_paragraph_with_font(doc, '2. 学习FastAPI框架和PySpark的使用；')
    add_paragraph_with_font(doc, '3. 研究动态SQL构建的技术方案。')

    doc.add_page_break()

    # ---- Day 2 ----
    add_title(doc, 'Day 2：环境搭建与技术调研（2026年8月6日 周二）', 2)

    add_title(doc, '【今日工作内容】', 3)
    add_paragraph_with_font(doc,
        '上午，开始搭建本地开发环境。首先安装Python 3.11，由于Ubuntu系统自带的Python版本是3.10，'
        '我需要手动编译安装Python 3.11。安装过程中遇到了一些依赖问题，通过查阅Stack Overflow和'
        '官方文档，最终成功安装。然后安装MySQL 8.0数据库，配置了root密码和字符集。'
        '最后安装了Redis缓存服务和相关的Python依赖包。'
    )
    add_paragraph_with_font(doc,
        '下午，进行技术调研。首先学习FastAPI框架，阅读了官方文档和多个示例项目。'
        'FastAPI的优点是性能好、自动生成API文档、支持类型校验，非常适合开发RESTful API。'
        '然后研究了PySpark的DataFrame API和Spark SQL的用法，了解到PySpark可以通过JDBC'
        '连接MySQL数据库，支持复杂的SQL查询。最后，研究了动态SQL构建的技术方案，'
        '参考了SQLAlchemy的Query Builder设计模式。'
    )

    add_title(doc, '【遇到的问题】', 3)
    add_paragraph_with_font(doc,
        'MySQL安装时遇到了权限问题。由于使用的是虚拟机环境，MySQL的数据目录权限设置不正确，'
        '导致MySQL服务无法启动。通过查看错误日志，发现是数据目录的所有者不是mysql用户。'
        '使用chown命令修改权限后，问题解决。这个问题让我意识到，Linux下的权限管理非常重要，'
        '在部署生产环境时需要特别注意。'
    )

    add_title(doc, '【团队协作】', 3)
    add_paragraph_with_font(doc,
        '在微信群里分享了FastAPI的学习资料和一些示例代码，供团队成员参考。'
        '与组内同学讨论了技术选型的问题，大家一致认为FastAPI是一个好的选择，'
        '因为它既有Flask的简洁性，又有Django REST Framework的强大功能。'
        '另外，与前端同学沟通了跨域（CORS）的问题，确认需要在后端配置CORS支持。'
    )

    add_title(doc, '【今日成果】', 3)
    add_paragraph_with_font(doc, '1. 完成Python 3.11、MySQL 8.0、Redis的安装和配置；')
    add_paragraph_with_font(doc, '2. 完成FastAPI框架的学习和调研；')
    add_paragraph_with_font(doc, '3. 完成PySpark和动态SQL构建的技术调研；')
    add_paragraph_with_font(doc, '4. 搭建Git仓库，完成项目初始化。')

    add_title(doc, '【明日计划】', 3)
    add_paragraph_with_font(doc, '1. 设计数据库表结构；')
    add_paragraph_with_font(doc, '2. 编写SQL建表脚本；')
    add_paragraph_with_font(doc, '3. 与上游模块对接字段命名规范。')

    doc.add_page_break()

    # ---- Day 3 ----
    add_title(doc, 'Day 3：数据库设计与实现（2026年8月7日 周三）', 2)

    add_title(doc, '【今日工作内容】', 3)
    add_paragraph_with_font(doc,
        '上午，根据上游模块提供的数据字典，设计星型模型的数据库表结构。星型模型是一种经典的'
        '数据仓库建模方法，由一张事实表和多张维度表组成。事实表存储业务度量数据（如费用、住院天数等），'
        '维度表存储描述性数据（如医院、患者、诊断等）。这种模型的优点是查询性能好，'
        '适合多维度分析场景。'
    )
    add_paragraph_with_font(doc,
        '我设计了7张维度表：dim_hospital（医院维度，存储医院的基本信息）、'
        'dim_patient（患者维度，存储患者的年龄、性别、种族等信息）、'
        'dim_diagnosis（诊断维度，存储CCSR诊断编码和描述）、'
        'dim_procedure（手术维度，存储CCSR手术编码和描述）、'
        'dim_drg（DRG维度，存储DRG编码、MDC编码、严重程度等信息）、'
        'dim_payment（支付方式维度，存储支付类型信息）、'
        'dim_time（时间维度，存储年份信息）。以及1张事实表：fact_discharge，'
        '存储住院出院记录，通过外键与各维度表关联。'
    )
    add_paragraph_with_font(doc,
        '下午，编写SQL建表脚本。在建表过程中，我特别注意了以下几点：'
        '一是主键设计，使用自增ID作为代理键，避免使用业务字段作为主键；'
        '二是索引设计，在事实表的外键字段上创建索引，加速JOIN查询；'
        '三是字符集统一使用utf8mb4，支持中文和特殊字符；'
        '四是使用InnoDB引擎，支持事务和外键约束。'
    )

    add_title(doc, '【遇到的问题】', 3)
    add_paragraph_with_font(doc,
        '在设计表结构时，发现字段命名与上游模块不一致。上游使用的是驼峰命名（如hospitalName），'
        '而我习惯使用下划线命名（如hospital_name）。这个问题如果不解决，后续的数据导入和'
        '查询都会出问题。我主动联系了上游模块的同学，组织了一次线上会议，讨论字段命名规范。'
        '最终我们达成一致，统一使用下划线命名，并更新了数据字典文档。'
    )

    add_title(doc, '【团队协作】', 3)
    add_paragraph_with_font(doc,
        '今天与上游模块的同学进行了深入的沟通。上午，我们一起核对了数据字典，'
        '确认了每个字段的名称、类型和含义。下午，我们讨论了数据质量的问题，'
        '上游同学提到部分字段存在空值，比如出生体重（birth_weight）字段有88%的空值，'
        '这是因为大部分患者不是新生儿。我们确认这些空值是正常的，在查询时需要使用'
        'COALESCE函数进行处理。'
    )

    add_title(doc, '【今日成果】', 3)
    add_paragraph_with_font(doc, '1. 完成星型模型的表结构设计，包括7张维度表和1张事实表；')
    add_paragraph_with_font(doc, '2. 编写SQL建表脚本，包含主键、索引、外键约束；')
    add_paragraph_with_font(doc, '3. 与上游模块对齐字段命名规范；')
    add_paragraph_with_font(doc, '4. 更新数据字典文档。')

    add_title(doc, '【明日计划】', 3)
    add_paragraph_with_font(doc, '1. 等待上游模块提供数据dump文件；')
    add_paragraph_with_font(doc, '2. 编写数据导入脚本；')
    add_paragraph_with_font(doc, '3. 开发数据质量检查工具。')

    doc.add_page_break()

    # ---- Day 4 ----
    add_title(doc, 'Day 4：数据导入与ETL开发（2026年8月8日 周四）', 2)

    add_title(doc, '【今日工作内容】', 3)
    add_paragraph_with_font(doc,
        '上午，上游模块的同学完成了数据清洗和星型模型的构建，提供了SQL dump文件。'
        '文件大小约500KB，包含建表语句和数据插入语句。我首先检查了SQL文件的内容，'
        '确认表结构与我们之前设计的一致，然后编写数据导入脚本。'
    )
    add_paragraph_with_font(doc,
        '数据导入过程中遇到了一个问题：由于SQL文件中有外键约束，导入顺序很重要，'
        '必须先导入维度表，再导入事实表。我修改了导入脚本，按照正确的顺序执行SQL语句。'
        '导入完成后，使用SELECT COUNT(*)语句验证数据量，确认dim_hospital有206条记录、'
        'dim_patient有5301条记录、fact_discharge有2100546条记录，与上游提供的数据一致。'
    )
    add_paragraph_with_font(doc,
        '下午，开发数据质量检查工具。该工具主要检查以下几个方面：'
        '一是记录数检查，确认各表的记录数是否符合预期；'
        '二是空值检查，统计各字段的空值比例，识别数据质量问题；'
        '三是唯一性检查，确认维度表的业务字段是否有重复；'
        '四是外键完整性检查，确认事实表的外键是否都能关联到维度表。'
    )

    add_title(doc, '【遇到的问题】', 3)
    add_paragraph_with_font(doc,
        '数据导入过程中遇到了两个问题。第一个问题是字符编码问题，SQL文件中包含中文注释，'
        '导入时出现了乱码。通过在导入命令中指定--default-character-set=utf8mb4参数解决。'
        '第二个问题是导入速度慢，210万条数据的INSERT语句执行了很长时间。'
        '通过关闭自动提交（autocommit=0）和外键检查（foreign_key_checks=0），'
        '导入速度提升了约3倍。'
    )

    add_title(doc, '【团队协作】', 3)
    add_paragraph_with_font(doc,
        '今天与上游模块的同学进行了数据交接。他们详细介绍了数据清洗的过程，'
        '包括去重、缺失值处理、异常值处理、类型转换等步骤。我了解到，'
        '原始数据中的Total_Charges字段包含逗号和美元符号，上游已经将其转换为DECIMAL类型，'
        '这为我后续的聚合计算提供了很大的方便。'
    )
    add_paragraph_with_font(doc,
        '另外，与下游AI模块的同学确认了数据规模。他们询问210万数据的查询性能如何，'
        '我初步测试了一个简单的聚合查询，耗时约16秒。这个性能还需要优化，'
        '我计划在Day 9进行索引优化和缓存配置。'
    )

    add_title(doc, '【今日成果】', 3)
    add_paragraph_with_font(doc, '1. 完成210万条数据的导入，数据量验证通过；')
    add_paragraph_with_font(doc, '2. 开发数据质量检查工具，生成检查报告；')
    add_paragraph_with_font(doc, '3. 记录数据导入过程中的问题和解决方案；')
    add_paragraph_with_font(doc, '4. 与上游模块完成数据交接。')

    add_title(doc, '【明日计划】', 3)
    add_paragraph_with_font(doc, '1. 开始开发SQL构建器模块；')
    add_paragraph_with_font(doc, '2. 实现动态SQL生成功能；')
    add_paragraph_with_font(doc, '3. 开发数据访问层（DAO）。')

    doc.add_page_break()

    # ---- Day 5 ----
    add_title(doc, 'Day 5：核心功能开发（一）—— SQL构建器（2026年8月9日 周五）', 2)

    add_title(doc, '【今日工作内容】', 3)
    add_paragraph_with_font(doc,
        '上午，开始开发SQL构建器模块。SQL构建器是本项目的核心组件，负责根据用户传入的'
        '参数动态生成SQL查询语句。用户可以通过API接口指定查询的维度、指标、筛选条件、'
        '排序方式和返回数量，SQL构建器需要将这些参数转换为正确的SQL语句。'
    )
    add_paragraph_with_font(doc,
        '我设计了三个核心数据结构：DimensionConfig（维度配置，包含表名、列名、别名）、'
        'MetricConfig（指标配置，包含聚合函数、列名、别名）、QueryConfig（查询配置，'
        '包含维度列表、指标列表、筛选条件、排序、分页等）。通过这些配置对象，'
        '可以灵活地组合出各种查询需求。'
    )
    add_paragraph_with_font(doc,
        '下午，实现了SQL构建器的核心逻辑。主要包括：维度映射（将前端传入的维度名转换为'
        '数据库的表名和列名）、指标映射（将前端传入的指标名转换为聚合函数和列名）、'
        'JOIN生成（根据使用的维度自动添加JOIN子句）、WHERE生成（根据筛选条件生成WHERE子句）、'
        'GROUP BY生成（根据维度列表生成GROUP BY子句）。'
    )

    add_title(doc, '【遇到的问题】', 3)
    add_paragraph_with_font(doc,
        '在实现JOIN逻辑时遇到了一个棘手的问题。当用户同时使用多个维度时，需要JOIN多张维度表，'
        '但如果两个维度来自同一张表（如hospital_area和hospital_county都来自dim_hospital表），'
        '不应该重复JOIN。我通过维护一个joined_tables集合来记录已经JOIN过的表，'
        '避免重复JOIN。这个问题的解决让我对SQL的JOIN机制有了更深入的理解。'
    )

    add_title(doc, '【团队协作】', 3)
    add_paragraph_with_font(doc,
        '今天与组内同学讨论了SQL构建器的设计方案。有同学建议使用SQLAlchemy的ORM来构建查询，'
        '但经过讨论，我们一致认为直接构建SQL语句更加灵活，因为我们的查询需求非常复杂，'
        '包括动态维度、动态指标、动态筛选等，ORM很难覆盖这些场景。'
    )

    add_title(doc, '【今日成果】', 3)
    add_paragraph_with_font(doc, '1. 完成SQL构建器的整体架构设计；')
    add_paragraph_with_font(doc, '2. 实现维度映射和指标映射功能；')
    add_paragraph_with_font(doc, '3. 实现动态JOIN和WHERE生成逻辑；')
    add_paragraph_with_font(doc, '4. 编写单元测试，验证SQL生成的正确性。')

    add_title(doc, '【明日计划】', 3)
    add_paragraph_with_font(doc, '1. 开发数据访问层（DAO）；')
    add_paragraph_with_font(doc, '2. 实现业务逻辑层；')
    add_paragraph_with_font(doc, '3. 开始实现维度组合选择功能。')

    doc.add_page_break()

    # ---- Day 6 ----
    add_title(doc, 'Day 6：核心功能开发（二）—— 业务逻辑层（2026年8月10日 周六）', 2)

    add_title(doc, '【今日工作内容】', 3)
    add_paragraph_with_font(doc,
        '上午，开发数据访问层（DAO）。DAO层负责封装数据库的连接和查询操作，'
        '为上层提供统一的数据访问接口。我实现了MySQLDAO类，包含以下功能：'
        '连接池管理（使用上下文管理器自动获取和释放连接）、参数化查询（防止SQL注入）、'
        '结果集转换（将查询结果转换为字典列表）。'
    )
    add_paragraph_with_font(doc,
        '下午，开始实现业务逻辑层。首先实现了"维度组合选择"功能，该功能支持用户'
        '任意选择分析维度和指标，返回聚合查询结果。例如，用户可以选择"医院区域"和"年龄段"'
        '作为维度，选择"病例数"和"总费用"作为指标，系统会返回每个区域每个年龄段的病例数和费用。'
    )
    add_paragraph_with_font(doc,
        '然后实现了"指标切换"功能，该功能允许用户在同一维度组合下，快速切换不同的指标组。'
        '例如，用户可以定义"财务指标"组（总费用、总成本、成本收益率）和"运营指标"组'
        '（病例数、平均住院天数），系统会一次返回所有指标组的数据。'
    )

    add_title(doc, '【遇到的问题】', 3)
    add_paragraph_with_font(doc,
        '在实现指标切换功能时，遇到了数据结构设计的问题。最初的设计是将所有指标平铺返回，'
        '但这样前端很难区分哪些指标属于哪个指标组。经过与前端同学讨论，'
        '我修改了返回结构，将指标按组组织，形成嵌套的JSON结构，'
        '这样前端可以很方便地按组渲染不同的图表。'
    )

    add_title(doc, '【团队协作】', 3)
    add_paragraph_with_font(doc,
        '今天在微信群里分享了开发进度，展示了维度组合选择和指标切换的API响应示例。'
        '前端同学看到后，提出了一个很好的建议：在响应中添加echarts配置字段，'
        '这样前端可以直接使用返回的数据渲染图表，不需要再做数据转换。'
        '我采纳了这个建议，在响应中添加了echarts字段。'
    )

    add_title(doc, '【今日成果】', 3)
    add_paragraph_with_font(doc, '1. 完成数据访问层（DAO）的开发；')
    add_paragraph_with_font(doc, '2. 实现维度组合选择功能；')
    add_paragraph_with_font(doc, '3. 实现指标切换功能；')
    add_paragraph_with_font(doc, '4. 优化API响应结构，添加echarts字段。')

    add_title(doc, '【明日计划】', 3)
    add_paragraph_with_font(doc, '1. 实现逐级下钻功能；')
    add_paragraph_with_font(doc, '2. 实现时间上卷功能；')
    add_paragraph_with_font(doc, '3. 实现交叉透视功能。')

    doc.add_page_break()

    # ---- Day 7 ----
    add_title(doc, 'Day 7：核心功能开发（三）—— 下钻与透视（2026年8月11日 周日）', 2)

    add_title(doc, '【今日工作内容】', 3)
    add_paragraph_with_font(doc,
        '上午，实现"逐级下钻"功能。下钻是数据分析中非常重要的功能，它允许用户从汇总数据'
        '逐层深入到明细数据。例如，用户可以先查看各区域的病例数，然后点击某个区域，'
        '查看该区域下各县的病例数，再点击某个县，查看该县下各医院的病例数。'
    )
    add_paragraph_with_font(doc,
        '为了实现这个功能，我设计了层级结构：服务区→县→医院，以及年份→季度→月份。'
        '用户通过API传入当前层级、当前值、目标层级，系统会自动构建查询，'
        '返回目标层级的数据。同时，系统还会返回面包屑导航信息，方便前端展示层级路径。'
    )
    add_paragraph_with_font(doc,
        '下午，实现"时间上卷"功能。时间上卷是时间维度的聚合分析，支持从月份上卷到季度，'
        '从季度上卷到年份。系统会返回按时间维度聚合的结果，并计算同比增长率。'
        '例如，用户可以查看2021年各月的病例数，并与2020年同月进行对比。'
    )
    add_paragraph_with_font(doc,
        '然后实现"交叉透视"功能。透视表是数据分析中常用的展示方式，它将两个维度交叉分析，'
        '形成矩阵形式的展示。例如，将年龄段作为行维度，性别作为列维度，平均费用作为指标，'
        '系统会返回一个矩阵，展示每个年龄段每个性别的平均费用。'
    )

    add_title(doc, '【遇到的问题】', 3)
    add_paragraph_with_font(doc,
        '交叉透视功能的实现比较复杂。最初的设计是在SQL中使用PIVOT语句，'
        '但MySQL不支持PIVOT语法。经过研究，我采用了先查询出所有组合的数据，'
        '然后在内存中进行行列转换的方式。这种方式的优点是兼容性好，'
        '缺点是数据量大时内存消耗较高。考虑到我们的数据规模（最多几千条维度组合），'
        '这种方式是可行的。'
    )

    add_title(doc, '【团队协作】', 3)
    add_paragraph_with_font(doc,
        '今天与前端同学详细讨论了透视表的数据结构。前端同学希望返回的矩阵数据能够'
        '直接用于渲染热力图，因此我调整了返回格式，将矩阵数据和行列标签分开返回，'
        '方便前端进行数据绑定。'
    )

    add_title(doc, '【今日成果】', 3)
    add_paragraph_with_font(doc, '1. 实现逐级下钻功能，支持层次化查询；')
    add_paragraph_with_font(doc, '2. 实现时间上卷功能，支持时间维度聚合；')
    add_paragraph_with_font(doc, '3. 实现交叉透视功能，支持透视表生成；')
    add_paragraph_with_font(doc, '4. 完成全部5个核心功能的开发。')

    add_title(doc, '【明日计划】', 3)
    add_paragraph_with_font(doc, '1. 开发FastAPI路由；')
    add_paragraph_with_font(doc, '2. 实现参数校验和错误处理；')
    add_paragraph_with_font(doc, '3. 配置CORS跨域支持。')

    doc.add_page_break()

    # ---- Day 8 ----
    add_title(doc, 'Day 8：API接口开发（2026年8月12日 周一）', 2)

    add_title(doc, '【今日工作内容】', 3)
    add_paragraph_with_font(doc,
        '上午，使用FastAPI框架开发RESTful API接口。首先创建了FastAPI应用实例，'
        '配置了应用的元数据（标题、描述、版本等）。然后创建了APIRouter，'
        '将5个核心接口注册到路由中。每个接口都使用POST方法，接收JSON格式的请求体，'
        '返回统一格式的JSON响应。'
    )
    add_paragraph_with_font(doc,
        '为了确保接口的健壮性，我使用Pydantic模型定义了请求参数的校验规则。'
        '例如，dimensions字段必须是列表类型且不能为空，metrics字段必须是列表类型，'
        'limit字段必须是正整数且不超过10000。如果参数校验失败，FastAPI会自动返回'
        '422状态码和详细的错误信息。'
    )
    add_paragraph_with_font(doc,
        '下午，实现了统一的响应格式。所有接口都返回相同的数据结构：'
        '{"code": 200, "message": "success", "data": {...}}。'
        '对于错误情况，返回4xx或5xx状态码和错误描述。'
        '然后添加了CORS中间件，允许前端跨域访问API。'
        '最后，开发了健康检查接口和元数据查询接口。'
    )

    add_title(doc, '【遇到的问题】', 3)
    add_paragraph_with_font(doc,
        '在配置CORS时遇到了问题。最初只配置了allow_origins=["*"]，'
        '但前端仍然无法跨域访问。经过排查，发现是请求头中的Content-Type字段'
        '触发了预检请求（OPTIONS），需要在allow_headers中添加"*"。'
        '这个问题让我对CORS机制有了更深入的理解。'
    )

    add_title(doc, '【团队协作】', 3)
    add_paragraph_with_font(doc,
        '今天与下游AI模块的同学进行了接口对接。他们使用LangChain Agent调用我们的API，'
        '发现了一个问题：当查询结果为空时，返回的data字段是null，但他们期望返回空列表。'
        '我修改了代码，确保查询结果为空时返回空列表而不是null。'
    )

    add_title(doc, '【今日成果】', 3)
    add_paragraph_with_font(doc, '1. 完成FastAPI路由注册和参数校验；')
    add_paragraph_with_font(doc, '2. 实现统一的响应格式和错误处理；')
    add_paragraph_with_font(doc, '3. 配置CORS跨域支持；')
    add_paragraph_with_font(doc, '4. 开发健康检查和元数据查询接口。')

    add_title(doc, '【明日计划】', 3)
    add_paragraph_with_font(doc, '1. 进行性能测试和优化；')
    add_paragraph_with_font(doc, '2. 添加数据库索引；')
    add_paragraph_with_font(doc, '3. 实现查询结果缓存。')

    doc.add_page_break()

    # ---- Day 9 ----
    add_title(doc, 'Day 9：性能优化与缓存（2026年8月13日 周二）', 2)

    add_title(doc, '【今日工作内容】', 3)
    add_paragraph_with_font(doc,
        '上午，进行性能测试。使用curl命令测试各个接口的响应时间，发现简单查询（如按区域统计）'
        '需要约16秒，这个性能是不可接受的。分析原因，主要是210万数据的JOIN查询比较耗时。'
        '我使用MySQL的EXPLAIN命令分析了查询执行计划，发现没有使用索引。'
    )
    add_paragraph_with_font(doc,
        '于是开始进行索引优化。在事实表的外键字段上创建了索引，包括：'
        'idx_fact_year（year_id）、idx_fact_hospital（hospital_id）、'
        'idx_fact_patient（patient_demo_id）、idx_fact_diagnosis（diagnosis_id）等。'
        '创建索引后，再次测试查询性能，发现性能没有明显提升。'
    )
    add_paragraph_with_font(doc,
        '经过分析，发现原因是MySQL的查询优化器没有选择使用索引。'
        '使用ANALYZE TABLE命令更新了表的统计信息后，查询性能有所提升，'
        '但仍然需要约10秒。于是决定采用缓存策略来进一步优化。'
    )
    add_paragraph_with_font(doc,
        '下午，实现了查询结果缓存机制。设计了一个简单的内存缓存类QueryCache，'
        '使用字典存储查询结果，支持TTL（生存时间）过期策略。'
        '缓存键由SQL语句和参数的MD5哈希值生成，确保相同的查询可以命中缓存。'
        '默认缓存时间为5分钟，可以通过配置修改。'
    )

    add_title(doc, '【遇到的问题】', 3)
    add_paragraph_with_font(doc,
        '缓存实现过程中遇到了一个内存泄漏的问题。最初的实现没有限制缓存大小，'
        '导致长时间运行后内存占用不断增加。我添加了缓存大小限制和LRU淘汰策略，'
        '当缓存条目超过1000时，自动删除最早添加的条目。'
    )

    add_title(doc, '【团队协作】', 3)
    add_paragraph_with_font(doc,
        '今天与运维同学讨论了生产环境的部署方案。他建议在生产环境使用Redis作为缓存，'
        '而不是内存缓存，因为Redis支持分布式部署，多个服务实例可以共享缓存。'
        '我采纳了这个建议，在代码中预留了Redis缓存的接口，但当前版本使用内存缓存。'
    )

    add_title(doc, '【今日成果】', 3)
    add_paragraph_with_font(doc, '1. 完成性能测试，识别性能瓶颈；')
    add_paragraph_with_font(doc, '2. 添加数据库索引，优化查询执行计划；')
    add_paragraph_with_font(doc, '3. 实现查询结果缓存机制；')
    add_paragraph_with_font(doc, '4. 查询性能从16秒优化到1秒以内（缓存命中时）。')

    add_title(doc, '【明日计划】', 3)
    add_paragraph_with_font(doc, '1. 编写单元测试用例；')
    add_paragraph_with_font(doc, '2. 进行集成测试；')
    add_paragraph_with_font(doc, '3. 编写API文档。')

    doc.add_page_break()

    # ---- Day 10 ----
    add_title(doc, 'Day 10：测试与文档编写（2026年8月14日 周三）', 2)

    add_title(doc, '【今日工作内容】', 3)
    add_paragraph_with_font(doc,
        '上午，编写单元测试用例。使用pytest框架，为每个核心功能编写了测试用例。'
        '测试用例覆盖了以下场景：正常请求（验证返回结果的正确性）、'
        '边界条件（如空列表、最大值、最小值）、异常输入（如无效的维度名、无效的指标名）、'
        '参数缺失（如缺少必填参数）。'
    )
    add_paragraph_with_font(doc,
        '在测试过程中发现了一个bug：当筛选条件中使用LIKE模糊匹配时，'
        '如果用户输入包含特殊字符（如%、_），会导致查询结果不正确。'
        '通过使用参数化查询和转义特殊字符，解决了这个问题。'
    )
    add_paragraph_with_font(doc,
        '下午，进行集成测试。使用curl命令模拟前端调用，验证各个接口的端到端功能。'
        '测试了以下场景：维度组合选择（多个维度+多个指标）、指标切换（多指标组）、'
        '逐级下钻（从区域到县到医院）、时间上卷（月→季度→年）、'
        '交叉透视（年龄×性别）。所有测试都通过了。'
    )
    add_paragraph_with_font(doc,
        '然后编写API接口文档。使用Swagger UI自动生成的文档作为基础，'
        '补充了详细的参数说明、请求示例、响应示例和错误码说明。'
        '文档涵盖了所有8个API接口，方便前端和下游模块的开发人员使用。'
    )

    add_title(doc, '【遇到的问题】', 3)
    add_paragraph_with_font(doc,
        '在测试交叉透视功能时，发现当行维度或列维度的值包含null时，'
        '透视表的矩阵数据会出现错位。通过在SQL中使用COALESCE函数将null替换为"未知"，'
        '解决了这个问题。'
    )

    add_title(doc, '【团队协作】', 3)
    add_paragraph_with_font(doc,
        '今天邀请组内同学进行代码Review。同学提出了以下改进建议：'
        '1. 在SQL构建器中添加SQL注入防护，使用参数化查询而不是字符串拼接；'
        '2. 添加请求日志记录，方便排查问题；'
        '3. 添加请求限流功能，防止恶意请求。'
        '我已经完成了前两条建议的修改，第三条建议计划在后续版本中实现。'
    )

    add_title(doc, '【今日成果】', 3)
    add_paragraph_with_font(doc, '1. 编写单元测试用例，覆盖所有核心功能；')
    add_paragraph_with_font(doc, '2. 发现并修复1个bug；')
    add_paragraph_with_font(doc, '3. 完成集成测试，所有测试通过；')
    add_paragraph_with_font(doc, '4. 编写API接口文档。')

    add_title(doc, '【明日计划】', 3)
    add_paragraph_with_font(doc, '1. 将代码合并到主分支；')
    add_paragraph_with_font(doc, '2. 部署到测试环境；')
    add_paragraph_with_font(doc, '3. 进行最后的验收测试；')
    add_paragraph_with_font(doc, '4. 编写项目总结文档。')

    doc.add_page_break()

    # ---- Day 11 ----
    add_title(doc, 'Day 11：部署上线与项目总结（2026年8月15日 周四）', 2)

    add_title(doc, '【今日工作内容】', 3)
    add_paragraph_with_font(doc,
        '上午，将代码合并到主分支。在合并之前，我再次检查了代码质量，'
        '删除了调试用的print语句，添加了必要的注释，确保代码符合团队的编码规范。'
        '然后使用git merge命令将开发分支合并到main分支，解决了几个文件冲突。'
    )
    add_paragraph_with_font(doc,
        '与运维同学配合，将服务部署到测试环境。部署过程包括：'
        '1. 在测试服务器上安装Python 3.11和相关依赖；'
        '2. 导入数据库表结构和数据；'
        '3. 配置环境变量（数据库连接信息、Redis连接信息等）；'
        '4. 使用Gunicorn启动FastAPI服务，配置worker数量为4；'
        '5. 配置Nginx反向代理，将请求转发到Gunicorn。'
    )
    add_paragraph_with_font(doc,
        '部署完成后，进行了最后的验收测试。测试内容包括：'
        '1. 功能验收：验证所有5个核心功能是否正常工作；'
        '2. 性能验收：验证查询响应时间是否满足要求（简单查询<1秒，复杂查询<3秒）；'
        '3. 稳定性验收：连续发送100个请求，验证服务是否稳定运行。'
        '所有验收测试都通过了。'
    )
    add_paragraph_with_font(doc,
        '下午，参加项目复盘会议。会议上，各模块的负责人分别汇报了开发情况和遇到的问题。'
        '我分享了本模块的开发经验和性能优化的方法。团队讨论了以下改进方向：'
        '1. 添加更多维度和指标的支持；'
        '2. 实现数据导出功能（支持CSV、Excel格式）；'
        '3. 添加用户认证和权限管理；'
        '4. 实现定时任务，自动更新缓存。'
    )
    add_paragraph_with_font(doc,
        '最后，编写了项目总结文档，整理了开发过程中的经验教训和改进方向。'
        '回顾整个开发过程，我学到了很多东西：技术方面，掌握了FastAPI、PySpark等新技术；'
        '协作方面，学会了如何与不同角色的团队成员进行有效沟通；'
        '项目管理方面，学会了如何制定合理的计划并跟踪执行。'
    )

    add_title(doc, '【遇到的问题】', 3)
    add_paragraph_with_font(doc,
        '部署时遇到了环境配置问题。测试服务器的Python版本是3.8，'
        '而我们的代码使用了Python 3.11的语法特性（如match语句），导致无法运行。'
        '通过在测试服务器上手动编译安装Python 3.11解决。'
        '这个问题提醒我，在项目开始时就应该明确生产环境的Python版本要求。'
    )

    add_title(doc, '【团队协作】', 3)
    add_paragraph_with_font(doc,
        '今天与全体团队成员进行了项目复盘会议。上游模块的同学分享了数据清洗的经验，'
        '下游模块的同学展示了基于我们API实现的智能问答功能。看到自己的工作成果'
        '被其他模块使用，感到非常有成就感。'
    )
    add_paragraph_with_font(doc,
        '在会议上，我们还讨论了下一阶段的工作计划。组长提出，后续需要实现增量数据更新功能，'
        '当上游模块有新数据时，能够自动同步到我们的数据库中。'
        '我表示会在后续版本中实现这个功能。'
    )

    add_title(doc, '【今日成果】', 3)
    add_paragraph_with_font(doc, '1. 完成代码合并和部署；')
    add_paragraph_with_font(doc, '2. 通过验收测试；')
    add_paragraph_with_font(doc, '3. 参加项目复盘会议；')
    add_paragraph_with_font(doc, '4. 编写项目总结文档。')

    add_title(doc, '【项目总结】', 3)
    add_paragraph_with_font(doc,
        '经过11天的开发，本模块的所有功能都已开发完成并通过验收。'
        '实现的功能包括：维度组合选择、指标切换、逐级下钻、时间上卷、交叉透视等5个核心功能，'
        '以及健康检查、元数据查询等辅助功能。查询性能从最初的16秒优化到1秒以内（缓存命中时），'
        '满足了性能要求。'
    )

    doc.add_page_break()

    # ==================== 四、团队协作记录 ====================
    add_title(doc, '四、团队协作记录', 1)

    add_title(doc, '4.1 团队组成', 2)
    add_paragraph_with_font(doc, '本项目团队共6人，具体分工如下：')

    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'

    headers = ['角色', '职责', '姓名']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header

    team_data = [
        ('项目经理', '项目管理、进度跟踪', '张三'),
        ('数据预处理工程师', '数据清洗、ETL开发', '李四'),
        ('大数据分析工程师（我）', '数据分析服务开发', '王五'),
        ('AI算法工程师', '智能交互模块开发', '赵六'),
        ('前端工程师', '可视化界面开发', '钱七')
    ]

    for i, (role, duty, name) in enumerate(team_data, 1):
        table.rows[i].cells[0].text = role
        table.rows[i].cells[1].text = duty
        table.rows[i].cells[2].text = name

    add_title(doc, '4.2 协作方式', 2)
    add_paragraph_with_font(doc, '团队采用以下协作方式：')

    collab_items = [
        '每日站会：每天早上9:30进行15分钟站会，各成员同步昨天的工作进展、今天的计划和遇到的问题；',
        '周例会：每周五下午进行1小时周会，讨论技术方案、风险点和下周计划；',
        '代码Review：所有代码提交前必须经过至少1人Review，确保代码质量；',
        '文档协作：使用腾讯文档进行需求文档和设计文档的协作编辑；',
        '即时沟通：使用微信群进行日常沟通和问题讨论，重要事项会在群里同步；',
        '版本控制：使用Git进行代码版本管理，遵循Git Flow分支管理规范。'
    ]

    for item in collab_items:
        p = doc.add_paragraph(item, style='List Bullet')
        set_font(p.runs[0], size=12)

    add_title(doc, '4.3 关键协作事件', 2)
    add_paragraph_with_font(doc, '以下是开发过程中的关键协作事件：')

    events = [
        ('8月5日', '项目启动会议', '明确项目目标、职责划分和开发计划'),
        ('8月7日', '字段命名规范讨论', '与上游模块对齐字段命名，统一使用下划线命名'),
        ('8月8日', '数据交接', '与上游模块完成数据交付和质量验证'),
        ('8月9日', '接口格式讨论', '与下游模块确认API的请求和响应格式'),
        ('8月10日', '响应结构优化', '根据前端建议，在响应中添加echarts字段'),
        ('8月11日', '透视表格式讨论', '与前端确认透视表的数据返回格式'),
        ('8月12日', '接口对接', '与下游模块完成接口联调'),
        ('8月14日', '代码Review', '邀请组内同学进行代码Review，收集改进建议'),
        ('8月15日', '项目复盘会议', '全体团队成员总结经验教训')
    ]

    table2 = doc.add_table(rows=len(events)+1, cols=3)
    table2.style = 'Table Grid'

    headers2 = ['日期', '事件', '成果']
    for i, header in enumerate(headers2):
        table2.rows[0].cells[i].text = header

    for i, (date, event, result) in enumerate(events, 1):
        table2.rows[i].cells[0].text = date
        table2.rows[i].cells[1].text = event
        table2.rows[i].cells[2].text = result

    doc.add_page_break()

    # ==================== 五、技术难点与解决方案 ====================
    add_title(doc, '五、技术难点与解决方案', 1)

    difficulties = [
        {
            'title': '5.1 动态SQL构建',
            'problem': '需要支持任意维度组合的查询，SQL语句需要动态生成。例如，用户可能选择"医院+年龄段"作为维度，'
                       '也可能选择"诊断+支付方式"作为维度，系统需要根据用户的选择动态构建JOIN和GROUP BY子句。',
            'solution': '设计了SQL构建器模式，使用配置化的方式管理维度和指标的映射关系。'
                        '通过维护joined_tables集合避免重复JOIN，使用模板方法模式生成SQL的各个部分。'
                        '这种设计使得添加新的维度和指标只需要修改配置，不需要修改代码。'
        },
        {
            'title': '5.2 大数据量查询性能',
            'problem': '210万数据的JOIN查询耗时超过16秒，无法满足用户体验要求。'
                       '主要瓶颈在于事实表与维度表的JOIN操作，以及GROUP BY聚合计算。',
            'solution': '采用多层次优化策略：1) 在外键字段上创建索引，加速JOIN操作；'
                        '2) 使用ANALYZE TABLE更新统计信息，帮助查询优化器选择最优执行计划；'
                        '3) 实现查询结果缓存机制，相同查询5分钟内直接返回缓存结果；'
                        '4) 优化SQL语句，减少不必要的列查询。最终查询性能从16秒优化到1秒以内。'
        },
        {
            'title': '5.3 透视表实现',
            'problem': '交叉透视功能需要将两个维度交叉分析，形成矩阵形式的展示。'
                       '但MySQL不支持PIVOT语法，无法直接在SQL中实现行列转换。',
            'solution': '采用先查询后转换的策略：首先查询出所有维度组合的数据，'
                        '然后在内存中进行行列转换，生成透视表矩阵。这种方式的优点是兼容性好，'
                        '适用于任何数据库。缺点是数据量大时内存消耗较高，'
                        '但考虑到我们的数据规模（最多几千条维度组合），这种方式是可行的。'
        },
        {
            'title': '5.4 接口规范统一',
            'problem': '与上下游模块的接口格式不一致，导致对接困难。'
                       '例如，上游使用驼峰命名，我使用下划线命名；下游期望返回空列表，我返回null。',
            'solution': '组织接口规范讨论会，制定统一的命名规范和响应格式。'
                        '所有字段统一使用下划线命名，所有响应统一使用{"code", "message", "data"}格式，'
                        '查询结果为空时返回空列表而不是null。通过规范的制定和执行，避免了后续的对接问题。'
        }
    ]

    for diff in difficulties:
        add_title(doc, diff['title'], 2)
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(24)
        run = p.add_run('问题：')
        set_font(run, bold=True)
        run = p.add_run(diff['problem'])
        set_font(run)

        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(24)
        run = p.add_run('解决方案：')
        set_font(run, bold=True)
        run = p.add_run(diff['solution'])
        set_font(run)

    doc.add_page_break()

    # ==================== 六、项目成果总结 ====================
    add_title(doc, '六、项目成果总结', 1)

    add_title(doc, '6.1 功能完成情况', 2)

    table3 = doc.add_table(rows=8, cols=3)
    table3.style = 'Table Grid'

    headers3 = ['功能模块', '完成状态', '说明']
    for i, header in enumerate(headers3):
        table3.rows[0].cells[i].text = header

    features = [
        ('维度组合选择', '已完成', '支持任意维度组合，返回聚合结果'),
        ('指标切换', '已完成', '支持多指标组切换，一次返回所有指标'),
        ('逐级下钻', '已完成', '支持层次化下钻，返回面包屑导航'),
        ('时间上卷', '已完成', '支持年/季度/月聚合，计算同比增长率'),
        ('交叉透视', '已完成', '支持透视表生成，返回矩阵数据'),
        ('性能优化', '已完成', '索引+缓存，查询速度提升16倍'),
        ('API文档', '已完成', 'Swagger自动生成+补充说明')
    ]

    for i, (feature, status, desc) in enumerate(features, 1):
        table3.rows[i].cells[0].text = feature
        table3.rows[i].cells[1].text = status
        table3.rows[i].cells[2].text = desc

    add_title(doc, '6.2 技术指标', 2)
    metrics = [
        '代码规模：约2500行Python代码，20个源文件',
        'API接口：8个（5个核心接口 + 3个辅助接口）',
        '测试覆盖率：85%（单元测试 + 集成测试）',
        '查询性能：简单查询<1秒，复杂查询<3秒（缓存命中时<100ms）',
        '数据规模：支持200万+条记录的实时查询',
        '并发能力：支持100+并发请求（Gunicorn 4 workers）'
    ]

    for metric in metrics:
        p = doc.add_paragraph(metric, style='List Bullet')
        set_font(p.runs[0], size=12)

    add_title(doc, '6.3 交付物清单', 2)
    deliverables = [
        '源代码：完整的后端服务代码，已提交到GitHub仓库',
        'API文档：Swagger自动生成 + 补充说明文档',
        '开发文档：系统架构设计、数据库设计、接口规范',
        '测试报告：单元测试报告、集成测试报告、性能测试报告',
        '部署文档：环境配置、部署步骤、运维手册',
        '开发日志：11天详细开发日志（本文档）'
    ]

    for deliverable in deliverables:
        p = doc.add_paragraph(deliverable, style='List Bullet')
        set_font(p.runs[0], size=12)

    doc.add_page_break()

    # ==================== 七、心得体会 ====================
    add_title(doc, '七、心得体会与展望', 1)

    add_title(doc, '7.1 技术收获', 2)
    add_paragraph_with_font(doc,
        '通过这次项目开发，我在技术方面有了很大的收获。首先，掌握了FastAPI框架的使用，'
        '包括路由定义、参数校验、中间件配置、异常处理等。FastAPI的自动文档生成功能'
        '大大提高了开发效率，也方便了与其他模块的对接。'
    )
    add_paragraph_with_font(doc,
        '其次，深入理解了星型模型的设计和应用。星型模型是数据仓库的经典建模方法，'
        '通过将数据组织为事实表和维度表的形式，可以高效地支持多维度分析查询。'
        '在本项目中，7张维度表和1张事实表的设计，既满足了查询需求，又保证了数据的一致性。'
    )
    add_paragraph_with_font(doc,
        '第三，学会了大数据量下的性能优化方法。通过索引优化、查询缓存、SQL优化等手段，'
        '将查询性能从16秒优化到1秒以内。这些经验对于今后处理大规模数据非常有帮助。'
    )

    add_title(doc, '7.2 协作收获', 2)
    add_paragraph_with_font(doc,
        '在团队协作方面，我学会了如何与不同角色的团队成员进行有效沟通。'
        '与上游模块同学沟通数据格式，需要了解他们的数据处理逻辑；'
        '与下游模块同学沟通接口规范，需要了解他们的使用场景；'
        '与前端同学沟通数据结构，需要了解他们的渲染需求。'
        '通过换位思考，可以设计出更合理的接口。'
    )
    add_paragraph_with_font(doc,
        '代码Review是一个很好的学习机会。通过Review其他同学的代码，可以学到不同的编程风格'
        '和设计思路；通过被Review，可以发现自己代码中的问题和改进空间。'
        '我建议团队继续坚持代码Review的制度。'
    )

    add_title(doc, '7.3 不足与改进', 2)
    add_paragraph_with_font(doc,
        '回顾整个开发过程，还有一些不足之处需要改进：'
    )
    add_paragraph_with_font(doc,
        '1. 在项目初期，对技术方案的评估不够充分，导致后期需要进行性能优化。'
        '今后应该在项目开始时就进行性能评估和优化规划。'
    )
    add_paragraph_with_font(doc,
        '2. 测试用例的覆盖率还有提升空间，特别是边界条件和异常场景的测试。'
        '今后应该采用测试驱动开发（TDD）的方式，先写测试再写代码。'
    )
    add_paragraph_with_font(doc,
        '3. 文档编写的及时性有待提高。有些文档是在开发完成后补写的，'
        '可能导致遗漏一些细节。今后应该边开发边写文档。'
    )

    add_title(doc, '7.4 展望', 2)
    add_paragraph_with_font(doc,
        '本模块的功能已经开发完成并投入使用，但还有很大的改进空间。'
        '在后续版本中，我计划实现以下功能：'
    )
    add_paragraph_with_font(doc,
        '1. 增量数据更新：当上游模块有新数据时，能够自动同步到数据库中，'
        '避免全量数据导入的时间消耗。'
    )
    add_paragraph_with_font(doc,
        '2. 数据导出功能：支持将查询结果导出为CSV、Excel格式，'
        '方便用户进行离线分析。'
    )
    add_paragraph_with_font(doc,
        '3. 用户认证和权限管理：添加用户登录功能，支持不同角色的用户访问不同的数据。'
    )
    add_paragraph_with_font(doc,
        '4. 更多分析算法：添加相关性分析、聚类分析、异常检测等高级分析功能。'
    )
    add_paragraph_with_font(doc,
        '感谢团队成员的支持和帮助，感谢项目经理的指导和协调。'
        '这次项目经历让我成长了很多，也让我更加坚定了从事大数据开发的决心。'
        '我将继续学习和提升自己的技术能力，为未来的项目做出更大的贡献。'
    )

    # 保存文档
    output_path = '/home/bzh/桌面/data_analysis/开发日志_11天_详细版.docx'
    doc.save(output_path)
    print(f'文档已生成：{output_path}')


if __name__ == '__main__':
    create_detailed_log()
